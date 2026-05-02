import os
import json
import logging
import numpy as np
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
#  1. VISUAL PROOF GENERATION
# ─────────────────────────────────────────────────────────────────────────────

def generate_radar_chart(category_scores: dict, output_path: str) -> str:
    """
    Generates a professional radar (spider) chart to visualize QA Risk Categories.
    """
    logger.info("Generating Radar Chart for Risk Categories...")
    
    labels = ['Complexité de la Matrice', 'Historique (Fiabilité)', 'Vulnérabilité d\'Exécution']
    values = list(category_scores.values())
    
    # Mathematical trick to close the radar chart loop
    values += values[:1]
    
    # Calculate angles for the radar
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]
    
    # Setup the plot with a modern aesthetic
    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    
    # Modern BIOPHARM Blue/Teal color
    modern_blue = '#0ea5e9' 
    
    # Plot the lines and fill the area
    ax.plot(angles, values, color=modern_blue, linewidth=2, linestyle='solid')
    ax.fill(angles, values, color=modern_blue, alpha=0.25) # 25% transparency is the secret sauce
    
    # Fix the labels and styling
    ax.set_yticklabels([]) # Hide ugly circular numbers
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=12, fontweight='bold', color='#333333')
    
    # Soften the gridlines
    ax.grid(color='#e5e7eb', linestyle='--', linewidth=1)
    ax.spines['polar'].set_color('none') # Remove the harsh outer black circle
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, transparent=True) # transparent=True makes it blend into Word perfectly
    plt.close()
    
    logger.info(f"Radar chart saved successfully at: {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
#  2. TEMPLATE INJECTION
# ─────────────────────────────────────────────────────────────────────────────

def generate_qa_report(json_data: dict, template_path: str, output_path: str):
    """
    Renders the Word document template using the provided JSON context and inserts the radar chart.
    """
    logger.info(f"Loading document template: {template_path}")
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found at {template_path}")

    doc = DocxTemplate(template_path)
    
    # Generate the visual proof (Radar Chart)
    chart_path = "radar_chart_temp.png"
    generate_radar_chart(json_data.get("category_scores", {}), chart_path)

    # Flatten the JSON dictionary for easier Jinja template access
    context = {
        "exact_risk_index": json_data.get("exact_risk_index"),
        "risk_category": json_data.get("risk_category"),
        "ml_model_used": json_data.get("ml_model_used"),
        "validation_recommendation": json_data.get("validation_recommendation"),
        "ich_justification": json_data.get("ich_justification"),
        
        # Flatten linearity_plan
        "linearity_levels": json_data.get("linearity_plan", {}).get("levels"),
        "linearity_replicates": json_data.get("linearity_plan", {}).get("replicates"),
        "linearity_range_pct": json_data.get("linearity_plan", {}).get("range_pct"),
        "linearity_reduction_pct": json_data.get("linearity_plan", {}).get("reduction_pct"),
        "linearity_rationale": json_data.get("linearity_plan", {}).get("rationale"),
        
        # Flatten accuracy_plan
        "accuracy_levels": json_data.get("accuracy_plan", {}).get("levels"),
        "accuracy_replicates": json_data.get("accuracy_plan", {}).get("replicates"),
        "accuracy_reduction_pct": json_data.get("accuracy_plan", {}).get("reduction_pct"),
    }
    
    # Dynamically flatten category scores
    for k, v in json_data.get("category_scores", {}).items():
        context[k] = v

    # Add the radar chart as an InlineImage
    context["radar_chart"] = InlineImage(doc, image_descriptor=chart_path, width=Inches(5.0))

    logger.info("Rendering Word Document with Context and Charts...")
    try:
        doc.render(context)
        doc.save(output_path)
        logger.info(f"Final Validation Report saved successfully at: {output_path}")
    except Exception as e:
        logger.error(f"Failed to render document: {e}")
        raise
    finally:
        # Cleanup temporary image file
        if os.path.exists(chart_path):
            os.remove(chart_path)


# ─────────────────────────────────────────────────────────────────────────────
#  3. EXECUTION FLOW
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("="*60)
    print("  QA Validation Report Generator")
    print("="*60)
    
    # Example Output Payload from our ML Engine
    mock_payload = {
        "exact_risk_index": 19.82,
        "risk_category": "LOW",
        "category_scores": {
            "matrix_complexity_risk": 41.21,
            "historical_performance_risk": 10.0,
            "execution_vulnerability_risk": 13.22
        },
        "ml_model_used": "GradientBoosting",
        "validation_recommendation": "Linearity: Reduce to 4 levels. Accuracy: Reduce to 3 levels. Total laboratory saving: 10 injections.",
        "ich_justification": "ICH REGULATORY JUSTIFICATION: The historical linearity and robustness data demonstrate sufficient prior knowledge, permitting partial validation effort reduction per ICH Q14 Section 5.",
        "linearity_plan": {
            "levels": 4,
            "replicates": 2,
            "range_pct": "70-90-110-130%",
            "reduction_pct": 46.7,
            "rationale": "Moderate risk with adequate historical linearity data supports partial reduction."
        },
        "accuracy_plan": {
            "levels": 3,
            "replicates": 2,
            "reduction_pct": 33.3
        }
    }
    
    TEMPLATE_FILE = "biopharm_aqbd_template.docx"
    OUTPUT_FILE = "Final_Validation_Report.docx"
    
    # For demonstration, generate a dummy template if it doesn't exist
    if not os.path.exists(TEMPLATE_FILE):
        logger.warning(f"Template '{TEMPLATE_FILE}' not found. Generating a dummy template for demonstration...")
        try:
            import docx
            dummy_doc = docx.Document()
            dummy_doc.add_heading('AQbD Final Validation Report', 0)
            dummy_doc.add_heading('Risk Analysis Overview', level=1)
            dummy_doc.add_paragraph('Risk Category: {{ risk_category }} (Index: {{ exact_risk_index }})')
            dummy_doc.add_paragraph('Model Used: {{ ml_model_used }}')
            dummy_doc.add_heading('Visual Risk Profile', level=1)
            dummy_doc.add_paragraph('{{ radar_chart }}')
            dummy_doc.add_heading('Recommendations & Justification', level=1)
            dummy_doc.add_paragraph('{{ validation_recommendation }}')
            dummy_doc.add_paragraph('{{ ich_justification }}')
            dummy_doc.save(TEMPLATE_FILE)
            logger.info(f"Dummy template created at {TEMPLATE_FILE}")
        except Exception as e:
            logger.error(f"Could not generate dummy template: {e}")
            exit(1)

    try:
        generate_qa_report(mock_payload, TEMPLATE_FILE, OUTPUT_FILE)
        print(f"\n[SUCCESS] The formal report has been generated: {OUTPUT_FILE}")
    except Exception as e:
        print(f"\n[ERROR] Process failed: {e}")
