import os
import json
import logging
from typing import Optional, Literal, List
from pathlib import Path

from pydantic import BaseModel, Field
from dotenv import load_dotenv

# Third-party libraries for document parsing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

# Instructor and Gemini for LLM extraction
try:
    import instructor
    import google.generativeai as genai
except ImportError:
    instructor = None
    genai = None

# Load environment variables (like GEMINI_API_KEY)
load_dotenv()

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
#  TARGET SCHEMA
# ─────────────────────────────────────────────────────────────────────────────

class FMEAExtraction(BaseModel):
    maturite_methode: Optional[int] = Field(description="1=High maturity (validated >=3 times), 2=Transferred, 3=New method")
    complexite_matrice: Optional[int] = Field(description="1=Pure API, 2=Tablet/Capsule, 3=Complex Cream/Emulsion")
    disponibilite_donnees: Optional[int] = Field(description="1=High (>=3 datasets), 2=Partial, 3=None")
    criticite_reglementaire: Optional[int] = Field(description="1=In-process control, 2=Stability, 3=Commercial release")
    risque_patient: Optional[int] = Field(description="1=Wide therapeutic index, 2=Standard, 3=Narrow therapeutic index")

class SequentialExtraction(BaseModel):
    current_x: Optional[List[float]] = Field(description="Array of concentration levels (e.g., [80.0, 90.0, 100.0])")
    current_y: Optional[List[float]] = Field(description="Array of corresponding instrument responses/areas")

class BayesianExtraction(BaseModel):
    historical_slopes: Optional[List[float]] = Field(description="Array of historical regression slopes")
    historical_intercepts: Optional[List[float]] = Field(description="Array of historical regression intercepts")

class SPARKMasterExtraction(BaseModel):
    fmea_inputs: FMEAExtraction
    sequential_inputs: SequentialExtraction
    bayesian_inputs: BayesianExtraction

# ─────────────────────────────────────────────────────────────────────────────
#  UNIVERSAL FILE PARSER
# ─────────────────────────────────────────────────────────────────────────────

class DocumentParser:
    """
    Extracts text and tabular data from PDF, DOCX, and XLSX files.
    """
    
    @staticmethod
    def parse(file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif ext == '.docx':
            return DocumentParser._parse_docx(file_path)
        elif ext == '.doc':
            return DocumentParser._parse_doc(file_path)
        elif ext in ['.xlsx', '.xls', '.csv']:
            return DocumentParser._parse_excel_csv(file_path)
        else:
            # Fallback to standard text reading
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    @staticmethod
    def _parse_doc(file_path: str) -> str:
        try:
            import win32com.client
        except ImportError:
            raise ImportError("pywin32 is required for .doc extraction on Windows.")
            
        content = []
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(abs_path)
            content.append(doc.Content.Text)
            
            for table in doc.Tables:
                content.append("\n[TABLE]")
                try:
                    # Avoid row iteration for merged cells. Use Range.Text directly.
                    raw_table = table.Range.Text
                    # \x07 is cell delimiter, \r is row delimiter in Word COM
                    formatted_table = raw_table.replace('\x07\r', '\n').replace('\x07', ' | ')
                    content.append(formatted_table)
                except Exception as table_err:
                    logger.warning(f"Could not parse a table due to formatting: {table_err}")
                content.append("[/TABLE]\n")
                
            doc.Close()
            word.Quit()
        except Exception as e:
            logger.error(f"Error parsing DOC {file_path}: {e}")
            try:
                word.Quit()
            except:
                pass
                
        return "\n".join(content)

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        if pdfplumber is None:
            raise ImportError("pdfplumber is required for PDF extraction.")
            
        content = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Extract text
                    text = page.extract_text()
                    if text:
                        content.append(text)
                    
                    # Extract tables to preserve structural context for LLM
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            # Convert table to simple markdown representation
                            md_table = "\n".join([" | ".join(str(cell).strip().replace('\n', ' ') for cell in row if cell is not None) for row in table])
                            content.append(f"\n[TABLE]\n{md_table}\n[/TABLE]\n")
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            
        return "\n".join(content)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        if docx is None:
            raise ImportError("python-docx is required for DOCX extraction.")
            
        content = []
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
                    
            for table in doc.tables:
                content.append("\n[TABLE]")
                for row in table.rows:
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    content.append(" | ".join(row_data))
                content.append("[/TABLE]\n")
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            
        return "\n".join(content)

    @staticmethod
    def _parse_excel_csv(file_path: str) -> str:
        if pd is None:
            raise ImportError("pandas is required for Excel/CSV extraction.")
            
        try:
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
                return df.to_markdown(index=False)
            else:
                content = []
                xls = pd.ExcelFile(file_path)
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    content.append(f"\n### Sheet: {sheet_name}\n")
                    content.append(df.to_markdown(index=False))
                return "\n".join(content)
        except Exception as e:
            logger.error(f"Error parsing Excel/CSV {file_path}: {e}")
            return ""

# ─────────────────────────────────────────────────────────────────────────────
#  STRUCTURED LLM EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

class LLMExtractor:
    """
    Uses Google Gemini API wrapped with Instructor to extract structured
    chemical parameters from document text into the SPARKMasterExtraction schema.
    """
    def __init__(self, api_key: str = None, model: str = "models/gemini-flash-latest"):
        if instructor is None or genai is None:
            raise ImportError("instructor and google-generativeai are required. Install with: pip install instructor google-generativeai")
            
        # Retrieve key from arguments, environment, or directly from GEMINI_API.TXT
        self.api_key = api_key or os.getenv("GEMINI_API_KEY")
        
        # If not found in env, explicitly parse GEMINI_API.TXT
        if not self.api_key and os.path.exists("GEMINI_API.TXT"):
            try:
                import re
                with open("GEMINI_API.TXT", "r", encoding="utf-8") as f:
                    match = re.search(r"AIza[0-9A-Za-z-_]{35}", f.read())
                    if match:
                        self.api_key = match.group(0)
            except Exception as e:
                logger.warning(f"Could not read GEMINI_API.TXT: {e}")

        if not self.api_key:
            raise ValueError("GEMINI_API_KEY is not set. Please provide it, set it in .env, or ensure GEMINI_API.TXT exists.")
        
        # Configure Gemini
        genai.configure(api_key=self.api_key)
        
        # Instructor patched Gemini client
        self.client = instructor.from_gemini(
            client=genai.GenerativeModel(model_name=model),
            mode=instructor.Mode.GEMINI_JSON,
        )
        
        self.system_prompt = (
            "You are an expert Pharmaceutical Data Entry Specialist and Analytical Chemist. "
            "Your job is to meticulously extract specific HPLC validation parameters from unstructured reports, "
            "PDFs, and Excel sheets into a highly structured JSON format supporting the SPARK API methodology.\n\n"
            "CRITICAL INSTRUCTIONS:\n"
            "1. Scrutinize the provided text and tables closely to infer the 1-3 FMEA scores based on their descriptions.\n"
            "2. Meticulously extract the X/Y data points (concentrations vs responses/areas) from any tabular data for the Sequential inputs.\n"
            "3. Extract historical slopes and intercepts for Bayesian inputs if mentioned.\n"
            "4. If a parameter is NOT explicitly stated or cannot be confidently inferred from the text, you MUST return `null` (None) for that field.\n"
            "5. DO NOT guess, fabricate, or hallucinate numbers. Silence (null) is far better than incorrect data."
        )

    def extract(self, document_text: str) -> SPARKMasterExtraction:
        """
        Extracts the SPARK schema from the provided document text.
        """
        logger.info("Sending document text to Gemini LLM for SPARK extraction...")
        
        # Construct message content explicitly conveying the system prompt 
        # and the user data since Gemini handles prompts slightly differently.
        prompt_payload = f"{self.system_prompt}\n\n--- DOCUMENT TEXT ---\n{document_text}\n--- END DOCUMENT TEXT ---\n\nExtract the required parameters matching the exact schema."
        
        response = self.client.chat.completions.create(
            response_model=SPARKMasterExtraction,
            messages=[
                {"role": "user", "content": prompt_payload}
            ]
        )
        
        return response


# ─────────────────────────────────────────────────────────────────────────────
#  EXECUTION FLOW
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("="*60)
    print("  QbD Document Extraction Pipeline (Gemini Edition)")
    print("="*60)
    
    # Path to one of the provided validation documents
    sample_file = "RVA du Dosage comprimé à 50mgversion 02.doc"
    
    # Try to parse a real file if available; else, fallback to mock text
    if os.path.exists(sample_file):
        print(f"Parsing file: {sample_file}")
        parser = DocumentParser()
        document_content = parser.parse(sample_file)
        print(f"Extracted {len(document_content)} characters from document.")
    else:
        print(f"Warning: '{sample_file}' not found in current directory.")
        print("Using sample dummy text for demonstration purposes...")
        document_content = """
        Validation Report: Coated Tablet Formulation
        The method utilizes an isocratic elution with a total run time of 10 minutes.
        Sample preparation involves 1 step: direct dissolve.
        Linearity demonstrated an R2 of 0.9999 over the target range.
        Precision testing yielded an RSD of 0.42%.
        Mean recovery was calculated at 99.8%.
        System suitability: Resolution Rs = 2.8, Tailing factor = 1.1.
        Nominal target concentration is 1.0 mg/mL.
        API stability is considered moderately stable, scoring a 3 on our internal scale.
        """

    try:
        # Initialize the Gemini extractor (it will automatically look for GEMINI_API.TXT)
        extractor = LLMExtractor()
        result = extractor.extract(document_content)
        
        print("\nExtracted JSON Payload:")
        print(result.model_dump_json(indent=2))
    except Exception as e:
        print(f"\nError during extraction: {e}")
